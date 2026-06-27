#!/usr/bin/env python3
"""Converte MANUAL_DO_USUARIO.md para .docx estilizado com a identidade visual do app."""

import re, os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paleta (app Reserva Quadra) ──
ROXO_PRIMARY = RGBColor(0x6D, 0x28, 0xD9)
ROXO_SECONDARY = RGBColor(0x7C, 0x3A, 0xED)
CINZA_BG = RGBColor(0xF3, 0xF4, 0xF6)
TEXTO_ESCURO = RGBColor(0x11, 0x18, 0x27)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(PROJECT_ROOT, "MANUAL_DO_USUARIO.md")
LOGO_PATH = os.path.join(PROJECT_ROOT, "public", "Complexo.jpeg")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "MANUAL_DO_USUARIO.docx")


def set_cell_shading(cell, color_hex):
    """Aplica fundo colorido a uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    tcPr.append(shading)


def add_run(paragraph, text, bold=False, italic=False, font_size=None,
            color=None, font_name=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def add_formatted_paragraph(doc, text, bold=False, italic=False, font_size=11,
                            color=None, alignment=None, space_after=6,
                            space_before=0, font_name="Calibri"):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        add_run(p, text, bold=bold, italic=italic, font_size=font_size,
                color=color, font_name=font_name)
    return p


def parse_inline(text):
    """Converte marcadores inline **bold** e *italic* para tuplas (texto, bold, italic)."""
    parts = []
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*)"
    last_end = 0
    for m in re.finditer(pattern, text):
        if m.start() > last_end:
            parts.append((text[last_end:m.start()], False, False))
        if m.group(2):
            parts.append((m.group(2), True, False))
        else:
            parts.append((m.group(3), False, True))
        last_end = m.end()
    if last_end < len(text):
        parts.append((text[last_end:], False, False))
    return parts


def render_inline_paragraph(doc_or_table, parts, **kwargs):
    """Adiciona parágrafo com runs formatados, aceita doc ou tabela."""
    p = doc_or_table.add_paragraph()
    for key in ('alignment', 'space_after', 'space_before'):
        if key in kwargs:
            setattr(p.paragraph_format, key, kwargs[key] if not isinstance(kwargs[key], int) else Pt(kwargs[key]))
    for texto, bold, italic in parts:
        run = p.add_run(texto)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = kwargs.get('color', TEXTO_ESCURO)
        run.font.size = Pt(kwargs.get('font_size', 11))
        run.font.name = kwargs.get('font_name', "Calibri")
    return p


def parse_markdown_tables(lines, start_idx):
    """Parseia tabela markdown a partir de start_idx. Retorna (headers, rows, end_idx)."""
    header_line = lines[start_idx].strip()
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    idx = start_idx + 2
    rows = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        idx += 1
    return headers, rows, idx


def add_table(doc, headers, rows):
    """Adiciona tabela formatada."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, font_size=10, color=BRANCO, font_name="Calibri")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "6D28D9")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parts = parse_inline(cell_text)
            for texto, bold, italic in parts:
                run = p.add_run(texto)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = TEXTO_ESCURO
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F4F6")

    return table


def add_cover_page(doc):
    """Cria a capa do documento."""
    # Espaço superior
    for _ in range(3):
        add_formatted_paragraph(doc, '', font_size=11)

    # Logo (se existir)
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.0))

    # Título principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    add_run(p_title, "Manual do Usuário", bold=True, font_size=32,
            color=ROXO_PRIMARY, font_name="Calibri Light")

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(6)
    add_run(p_sub, "Reserva Quadra", bold=False, font_size=22,
            color=ROXO_SECONDARY, font_name="Calibri Light")

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_before = Pt(18)
    add_run(p_desc, "Sistema de agendamento da quadra poliesportiva", font_size=12,
            color=TEXTO_ESCURO, font_name="Calibri")
    p_desc2 = doc.add_paragraph()
    p_desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_desc2, "do Complexo Júlio Prestes", font_size=12,
            color=TEXTO_ESCURO, font_name="Calibri")

    # Data
    from datetime import date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(48)
    add_run(p_date, f"Gerado em {date.today().strftime('%d/%m/%Y')}", font_size=11,
            color=RGBColor(0x9C, 0xA3, 0xAF), font_name="Calibri")

    # Quebra de página
    doc.add_page_break()


def add_toc(doc):
    """Adiciona sumário e campo TOC automático do Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_run(p, "Sumário", bold=True, font_size=20, color=ROXO_PRIMARY,
            font_name="Calibri Light")

    # Campo TOC
    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.space_after = Pt(12)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-4" \\h \\z \\u </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_toc = p_toc.add_run()
    run_toc._r.append(fldChar1)
    run_toc._r.append(instrText)
    run_toc._r.append(fldChar2)
    run_toc._r.append(fldChar3)

    doc.add_page_break()


def add_header_footer(doc):
    """Adiciona cabeçalho e rodapé com numeração de páginas."""
    for section in doc.sections:
        # Cabeçalho
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("Manual do Usuário — Reserva Quadra")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run.font.name = "Calibri"

        # Rodapé com página
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "Página "
        run1 = fp.add_run("Página ")
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        # Campo PAGE
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instrText_page = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

        run_page = fp.add_run()
        run_page.font.size = Pt(8)
        run_page.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_page._r.append(fldChar_begin)
        run_page._r.append(instrText_page)
        run_page._r.append(fldChar_sep)
        run_page._r.append(fldChar_end)


def convert_md_to_docx():
    """Função principal."""
    if not os.path.exists(MD_PATH):
        print(f"ERRO: {MD_PATH} não encontrado.")
        sys.exit(1)

    with open(MD_PATH, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()

    doc = Document()
    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Capa + Sumário
    add_cover_page(doc)
    add_toc(doc)

    # Processar linhas
    i = 0
    in_code_block = False
    code_buffer = []
    while i < len(md_lines):
        line = md_lines[i].rstrip()

        # Code block
        if line.startswith('```'):
            if in_code_block:
                # Fecha code block
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Cm(0.5)
                shading_para = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F3F4F6" w:val="clear"/>'
                )
                pPr = p._p.get_or_add_pPr()
                pPr.append(shading_para)
                for c_line in code_buffer:
                    run = p.add_run(c_line + '\n')
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = TEXTO_ESCURO
                code_buffer = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Separador
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Linha horizontal com borda inferior
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Tabela
        if line.startswith('|') and i + 1 < len(md_lines) and '---' in md_lines[i + 1]:
            headers, rows, end = parse_markdown_tables(md_lines, i)
            if headers:
                add_formatted_paragraph(doc, '', font_size=4, space_after=2)
                add_table(doc, headers, rows)
                add_formatted_paragraph(doc, '', font_size=4, space_after=6)
            i = end
            continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Barra lateral roxa
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="6D28D9"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            # Conteúdo inline
            parts = parse_inline(text)
            for texto, bold, italic in parts:
                run = p.add_run(texto)
                run.bold = bold
                run.italic = True
                run.font.size = Pt(11)
                run.font.color.rgb = ROXO_SECONDARY
                run.font.name = "Calibri"
            i += 1
            continue

        # Headings
        h_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            # Remove links [texto](#...)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)

            sizes = {1: 24, 2: 18, 3: 14, 4: 12}
            size = sizes.get(level, 11)

            # H1 - centralizado
            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                # Linha decorativa abaixo
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="12" w:space="4" w:color="6D28D9"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                add_run(p, text, bold=True, font_size=size, color=ROXO_PRIMARY,
                        font_name="Calibri Light")
                i += 1
                continue

            # H2/H3/H4
            space_before = {2: 20, 3: 14, 4: 10}.get(level, 10)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(6)
            bold = level <= 3
            italic = level == 4
            add_run(p, text, bold=bold, italic=italic, font_size=size,
                    color=ROXO_PRIMARY, font_name="Calibri" if level >= 3 else "Calibri Light")
            i += 1
            continue

        # Parágrafo normal
        if line.strip():
            parts = parse_inline(line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(2)
            for texto, bold, italic in parts:
                run = p.add_run(texto)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(11)
                run.font.name = "Calibri"
                run.font.color.rgb = TEXTO_ESCURO
            i += 1
            continue

        # Linha vazia
        i += 1

    # Cabeçalho e rodapé
    add_header_footer(doc)

    # Salvar
    doc.save(OUTPUT_PATH)
    print(f"Documento gerado: {OUTPUT_PATH}")
    print(f"Tamanho: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == '__main__':
    convert_md_to_docx()
